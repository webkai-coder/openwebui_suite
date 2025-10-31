import base64
import binascii
from collections import OrderedDict
from io import BytesIO
from typing import Iterable, List, Literal, Optional, Tuple

import re
import spacy
from docx import Document
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

router = APIRouter()
nlp = spacy.load("de_core_news_sm")

EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
ENTITY_LABELS = ("PER", "ORG", "LOC")


class ScrubRequest(BaseModel):
    mode: Literal["text", "docx"] = "text"
    text: Optional[str] = None
    docx_base64: Optional[str] = None
    return_docx: bool = False
    preserve_formatting: bool = True
    include_clean_text: bool = False


class ScrubResponse(BaseModel):
    clean_text: Optional[str] = None
    clean_docx_base64: Optional[str] = None
    warnings: List[str] = Field(default_factory=list)


@router.post("/scrub", response_model=ScrubResponse)
def scrub_text(item: ScrubRequest) -> ScrubResponse:
    warnings: List[str] = []

    if item.mode == "text":
        if not item.text:
            raise HTTPException(status_code=400, detail="Field 'text' is required for mode='text'.")
        clean_text, _ = _scrub_text(item.text)
        return ScrubResponse(clean_text=clean_text, warnings=warnings)

    if not item.docx_base64:
        raise HTTPException(status_code=400, detail="Field 'docx_base64' is required for mode='docx'.")

    try:
        doc_bytes = base64.b64decode(item.docx_base64)
    except (binascii.Error, ValueError) as exc:
        raise HTTPException(status_code=400, detail="Field 'docx_base64' must be valid base64.") from exc

    original_doc = Document(BytesIO(doc_bytes))
    plain_text = _collect_document_text(original_doc)
    clean_text, replacements = _scrub_text(plain_text)

    clean_docx_base64: Optional[str] = None
    if item.return_docx:
        redacted_doc = Document(BytesIO(doc_bytes))
        warnings.extend(_apply_replacements_to_document(redacted_doc, replacements, item.preserve_formatting))
        buffer = BytesIO()
        redacted_doc.save(buffer)
        clean_docx_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")

    response_text = clean_text if item.include_clean_text else None
    return ScrubResponse(clean_text=response_text, clean_docx_base64=clean_docx_base64, warnings=warnings)


def _scrub_text(text: str) -> Tuple[str, List[Tuple[str, str]]]:
    segments: List[Tuple[int, int, str, str]] = []
    placeholder_map: OrderedDict[str, str] = OrderedDict()

    email_counter = 0
    for match in EMAIL_PATTERN.finditer(text):
        raw = match.group(0)
        if raw not in placeholder_map:
            email_counter += 1
            placeholder_map[raw] = f"<EMAIL_{email_counter}>"
        segments.append((match.start(), match.end(), placeholder_map[raw], raw))

    label_counters = {label: 0 for label in ENTITY_LABELS}
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ not in ENTITY_LABELS:
            continue
        raw = text[ent.start_char:ent.end_char]
        if raw not in placeholder_map:
            label_counters[ent.label_] += 1
            placeholder_map[raw] = f"<{ent.label_}_{label_counters[ent.label_]}>"
        segments.append((ent.start_char, ent.end_char, placeholder_map[raw], raw))

    segments.sort(key=lambda item: (item[0], -(item[1] - item[0])))

    cleaned_parts: List[str] = []
    last_index = 0
    ordered_replacements: List[Tuple[str, str]] = []
    seen: set[str] = set()

    for start, end, placeholder, raw in segments:
        if start < last_index:
            continue
        cleaned_parts.append(text[last_index:start])
        cleaned_parts.append(placeholder)
        if raw not in seen:
            ordered_replacements.append((raw, placeholder))
            seen.add(raw)
        last_index = end

    cleaned_parts.append(text[last_index:])
    return "".join(cleaned_parts), ordered_replacements


def _collect_document_text(document: Document) -> str:
    chunks: List[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            chunks.append(paragraph.text)
    for table in document.tables:
        chunks.extend(_collect_table_text(table))
    return "\n".join(chunks)


def _collect_table_text(table) -> List[str]:
    chunks: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    chunks.append(paragraph.text)
            for nested in cell.tables:
                chunks.extend(_collect_table_text(nested))
    return chunks


def _apply_replacements_to_document(
    document: Document,
    replacements: List[Tuple[str, str]],
    preserve_formatting: bool,
) -> List[str]:
    ordered_replacements = sorted(replacements, key=lambda pair: len(pair[0]), reverse=True)
    warnings: List[str] = []
    warnings.extend(_replace_in_paragraph_sequence(document.paragraphs, ordered_replacements, preserve_formatting))
    for table in document.tables:
        warnings.extend(_replace_in_table(table, ordered_replacements, preserve_formatting))
    return warnings


def _replace_in_table(table, replacements: List[Tuple[str, str]], preserve_formatting: bool) -> List[str]:
    warnings: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            warnings.extend(_replace_in_paragraph_sequence(cell.paragraphs, replacements, preserve_formatting))
            for nested in cell.tables:
                warnings.extend(_replace_in_table(nested, replacements, preserve_formatting))
    return warnings


def _replace_in_paragraph_sequence(paragraphs, replacements: List[Tuple[str, str]], preserve_formatting: bool) -> List[str]:
    warnings: List[str] = []
    for paragraph in paragraphs:
        warnings.extend(_replace_in_paragraph(paragraph, replacements, preserve_formatting))
    return warnings


def _replace_in_paragraph(paragraph, replacements: List[Tuple[str, str]], preserve_formatting: bool) -> List[str]:
    warnings: List[str] = []
    desired_text = _apply_replacements(paragraph.text, replacements)

    if preserve_formatting:
        for run in paragraph.runs:
            run.text = _apply_replacements(run.text, replacements)
        if paragraph.text != desired_text:
            _rewrite_paragraph(paragraph, desired_text)
            warnings.append("Absatz musste für vollständige Redaktion neu gesetzt werden.")
    else:
        _rewrite_paragraph(paragraph, desired_text)

    return warnings


def _apply_replacements(value: str, replacements: Iterable[Tuple[str, str]]) -> str:
    result = value
    for original, redacted in replacements:
        result = result.replace(original, redacted)
    return result


def _rewrite_paragraph(paragraph, text: str) -> None:
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
    if text:
        paragraph.add_run(text)


tool_spec = {
    "name": "scrub",
    "description": (
        "Entfernt PII (E-Mail, Namen, Orte) aus Texten oder DOCX-Dateien. "
        "Optional wird das Ergebnis wieder als DOCX geliefert."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "mode": {
                "type": "string",
                "enum": ["text", "docx"],
                "description": "Verarbeitungsmodus: Rohtext oder DOCX-Datei."
            },
            "text": {
                "type": "string",
                "description": "Nur für mode=text: Zu bereinigender Text."
            },
            "docx_base64": {
                "type": "string",
                "description": "Nur für mode=docx: Base64-kodierte DOCX-Datei."
            },
            "return_docx": {
                "type": "boolean",
                "default": False,
                "description": "Nur für mode=docx: True liefert ein redigiertes DOCX ."
            },
            "preserve_formatting": {
                "type": "boolean",
                "default": True,
                "description": "Formatierung im DOCX nach Möglichkeit behalten."
            },
            "include_clean_text": {
                "type": "boolean",
                "default": False,
                "description": "Bereinigten Volltext zusätzlich zurückgeben?"
            }
        },
        "required": ["mode"],
        "additionalProperties": False
    },
    "endpoint": "/scrub",
    "method": "POST",
    "output_key": "clean_text",
    "output_schema": {
        "type": "object",
        "properties": {
            "clean_text": {
                "type": ["string", "null"],
                "description": "Bereinigter Text (falls angefordert/verfügbar)."
            },
            "clean_docx_base64": {
                "type": ["string", "null"],
                "description": "Bereinigtes DOCX, base64-kodiert (falls angefordert)."
            },
            "warnings": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Hinweise, z. B. wenn Formatierung neu gesetzt wurde."
            }
        },
        "required": ["warnings"],
        "additionalProperties": False
    }
}
